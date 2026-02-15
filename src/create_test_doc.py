"""Generate a test Word document with various OECD Style Guide violations."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_test_document(output_path: str = "../data/test_style_errors.docx"):
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────
    title = doc.add_heading("OECD Policy Brief on Digital Transformation", level=0)

    # ──────────────────────────────────────────────────────────────────
    # ERROR 1: Acronym not spelled out on first use (violates rule_26)
    # Should be: "The Organisation for Economic Co-operation and Development (OECD)"
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    p = doc.add_paragraph(
        "The OECD has published several reports on digital transformation. "
        "GDP growth in OECD countries averaged 2.3% in 2024. "
        "The IMF and WTO also contributed to the analysis. "
        "AI adoption is accelerating across all sectors."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 2: Capitalisation errors (violates rule_59, rule_60)
    # "member countries" should be "Member countries"
    # "the organisation" should be "the Organisation" when referring to OECD
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "The organisation has been working with its member countries to develop "
        "guidelines for responsible AI governance. Several non-member economies "
        "have also expressed interest in these initiatives. The oecd council met "
        "in January to discuss the Digital government strategy."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 3: Wrong number formatting (violates rule_14 / numbers rules)
    # Numbers below 10 should be written out; large numbers need spacing
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("2. Key Findings", level=1)
    p = doc.add_paragraph(
        "The study covered 3 countries and surveyed 5 institutions over a "
        "period of 2 years. Approximately 1500000 data points were collected. "
        "The results showed that 8 out of 10 respondents supported the policy, "
        "representing sixty-seven % of the total sample."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 4: Date and time formatting issues (violates rule_10)
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "The project ran from June 3rd, 2024 to December 1st, 2024. "
        "Data was collected during the Spring of 2024. "
        "The final report was submitted on the 15th of November 2024. "
        "Results are expected in the 2nd quarter of 2025."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 5: Hyphenation errors (violates rule_11)
    # "e-words" like "e-Government" should be "e-government" (lowercase)
    # "decision making" should be "decision-making" when used as adjective
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("3. Policy Recommendations", level=1)
    p = doc.add_paragraph(
        "The E-Government initiative has led to significant improvements in "
        "service delivery. The decision making process should be streamlined. "
        "Cross border cooperation remains essential for long term growth. "
        "The M-Commerce sector has seen year on year increases."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 6: Inclusive language issues (violates rule_12)
    # Use gender-neutral language; avoid "he" as generic pronoun
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "When a researcher submits his findings, he should ensure that all "
        "data sources are properly cited. The chairman of the committee "
        "presented the findings to the manpower division. Each country "
        "should submit his annual report by the deadline."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 7: Citation/reference issues (violates rule_34, rule_39)
    # Missing sources, wrong citation format
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("4. Data Analysis", level=1)
    p = doc.add_paragraph(
        "According to recent studies, digital transformation has increased "
        "productivity by 15%. Research shows that AI adoption varies significantly "
        "across countries. The data confirms that investment in digital infrastructure "
        "correlates strongly with economic growth (see various sources)."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 8: Figure/table source missing (violates rule_51)
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "Figure 1. Digital adoption index across OECD countries, 2020-2024"
    )
    p.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else None
    # Table without source
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "Country"
    headers[1].text = "Digital Index 2020"
    headers[2].text = "Digital Index 2024"
    data_rows = [
        ("France", "67.2", "78.5"),
        ("Germany", "71.8", "82.1"),
        ("Japan", "65.4", "76.9"),
    ]
    for i, (country, idx2020, idx2024) in enumerate(data_rows):
        row = table.rows[i + 1].cells
        row[0].text = country
        row[1].text = idx2020
        row[2].text = idx2024

    # Missing: "Source: OECD (2024), Digital Government Index..."
    doc.add_paragraph("")  # No source line!

    # ──────────────────────────────────────────────────────────────────
    # ERROR 9: Punctuation and spacing issues (violates rule_16)
    # Wrong use of ampersand, missing Oxford comma, etc.
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("5. Conclusions & Next Steps", level=1)  # Ampersand in heading!
    p = doc.add_paragraph(
        "The OECD,the EU and the World Bank have collaborated on this project. "  # missing space after comma
        "Countries should focus on: infrastructure , skills, and governance. "  # extra space before comma
        "The results are clear ...digital transformation is essential. "  # wrong ellipsis
        "For more information,contact the OECD Secretariat."  # missing space
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 10: Spelling - American vs British English (violates rule_18)
    # OECD uses British spelling
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "The organization should prioritize the digitalization of public services. "  # organization → organisation, prioritize → prioritise
        "This program aims to analyze labor market trends and optimize "  # programme, analyse, labour, optimise
        "resource utilization across all centers of operation."  # utilisation, centres
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 11: Wrong formatting - bold/italic misuse (violates rule_13)
    # Underlining should not be used; italic for specific purposes only
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("6. References", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.underline = True  # ERROR: Underlining should not be used
    run = p.add_run("The following publications are ")
    run = p.add_run("highly recommended")
    run.underline = True  # ERROR: underline
    run = p.add_run(" for further reading.")

    # ──────────────────────────────────────────────────────────────────
    # ERROR 12: Reference formatting errors (violates rule_40)
    # Wrong author name format, missing elements
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "John Smith (2024), Digital Government Review, OECD Publishing, Paris."
        # Should be: Smith, J. (2024), ...
    )
    p = doc.add_paragraph(
        'OECD, "AI Policy Framework", 2023, Paris.'
        # Should be: OECD (2023), "AI Policy Framework", OECD Publishing, Paris.
    )
    p = doc.add_paragraph(
        "See https://www.oecd.org/digital for more details."
        # Bare URL without proper formatting
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 13: Country name issues (violates rule_9 / OECD naming conventions)
    # ──────────────────────────────────────────────────────────────────
    doc.add_heading("Annex A. Country Coverage", level=1)
    p = doc.add_paragraph(
        "This report covers the following economies: "
        "South Korea, Czech Republic, Turkey, and England. "
        # Should be: Korea, Czechia, Türkiye, United Kingdom
        "Data was also collected from Mainland China and the occupied Palestinian territories."
    )

    # ──────────────────────────────────────────────────────────────────
    # ERROR 14: Referring to the OECD incorrectly (violates rule_15)
    # ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "The OECD organization was founded in 1961. "  # Don't say "OECD organization"
        "We believe this policy will benefit all our members. "  # Avoid first person, "our members"
        "The OECD's headquarters are located in Paris, France."
    )

    # Save
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Test document saved to: {output.resolve()}")
    print(f"\nErrors included:")
    print(f"  1. Acronyms not spelled out on first use (OECD, GDP, IMF, WTO, AI)")
    print(f"  2. Capitalisation: 'member countries', 'the organisation', 'oecd council'")
    print(f"  3. Number formatting: '3 countries', '1500000', 'sixty-seven %'")
    print(f"  4. Date formatting: '3rd', '1st', 'Spring' capitalised, '15th of November'")
    print(f"  5. Hyphenation: 'E-Government', 'decision making', 'Cross border', 'long term'")
    print(f"  6. Inclusive language: 'his findings', 'chairman', 'manpower'")
    print(f"  7. Missing citations: 'recent studies', 'Research shows', 'see various sources'")
    print(f"  8. Table/figure without source attribution")
    print(f"  9. Punctuation: ampersand in heading, spacing errors, wrong ellipsis")
    print(f" 10. American spelling: 'organization', 'prioritize', 'analyze', 'labor'")
    print(f" 11. Underlining (forbidden), formatting misuse")
    print(f" 12. Reference format: wrong author name order, missing elements")
    print(f" 13. Country names: 'South Korea', 'Czech Republic', 'Turkey', 'England'")
    print(f" 14. OECD references: 'OECD organization', first person 'we/our'")


if __name__ == "__main__":
    create_test_document()
