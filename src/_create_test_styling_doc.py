"""
Create a test Word document with OpenXML-level styling violations.

These violations are NOT visible when reading the text, but are present in the
underlying XML formatting runs — exactly the kind of thing our quality checker
(which feeds run-level formatting to the LLM) should catch.

Targeted OECD Style Guide rules:
  rule_138 (mandatory): Do not apply bold/italics/underlining to create hierarchy
  rule_141 (mandatory): Don't overuse italics for emphasis
  rule_149 (mandatory): Use bold very sparingly, only if essential for meaning
  rule_150 (mandatory): Don't apply bold as a formatting choice for hierarchy
  rule_152 (mandatory): Don't use underlining
  rule_177 (mandatory): Don't italicise or put in bold units of measure
  rule_140 (mandatory): Italicise titles of reports, books, journals
  rule_143 (mandatory): Use roman type for section/chapter titles
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

OUT = Path(__file__).resolve().parent.parent / "data" / "test_styling_violations.docx"


def main() -> None:
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────
    doc.add_heading("Climate Action and Economic Growth", level=1)

    # ──────────────────────────────────────────────────────────────────
    # 1. CLEAN paragraph (no violations) — baseline
    # ──────────────────────────────────────────────────────────────────
    p1 = doc.add_paragraph()
    p1.add_run(
        "The OECD has long advocated for policies that align environmental "
        "sustainability with economic development. This report examines the "
        "trade-offs and synergies between climate action and GDP growth across "
        "member countries."
    )

    # ──────────────────────────────────────────────────────────────────
    # 2. BOLD body text — violates rule_149 and rule_150
    #    Text reads normally, but entire paragraph is bold in OpenXML.
    # ──────────────────────────────────────────────────────────────────
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Carbon pricing mechanisms have been adopted by 46 national and "
        "32 subnational jurisdictions, covering approximately 23% of global "
        "greenhouse gas emissions."
    )
    run2.bold = True  # ← OpenXML violation: body text should NOT be bold

    # ──────────────────────────────────────────────────────────────────
    # 3. UNDERLINED body text — violates rule_152
    #    Underlining is prohibited except in auto-generated URLs.
    # ──────────────────────────────────────────────────────────────────
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        "Subsidies for fossil fuels remain a significant barrier to "
        "achieving net-zero targets by 2050."
    )
    run3.underline = WD_UNDERLINE.SINGLE  # ← OpenXML violation: no underline

    # ──────────────────────────────────────────────────────────────────
    # 4. ITALIC on common English words — violates rule_141
    #    Italics should NOT be used on established English words.
    # ──────────────────────────────────────────────────────────────────
    p4 = doc.add_paragraph()
    p4.add_run("The concept of ")
    run4a = p4.add_run("sustainability")
    run4a.italic = True  # ← violation: common English word, not foreign term
    p4.add_run(
        " is central to modern economic policy. Governments must balance "
        "short-term fiscal pressures with long-term environmental goals."
    )

    # ──────────────────────────────────────────────────────────────────
    # 5. Book title NOT in italics — violates rule_140
    #    Book/report titles must be italicised.
    # ──────────────────────────────────────────────────────────────────
    p5 = doc.add_paragraph()
    p5.add_run(
        "As noted in the OECD Economic Outlook 2024, fiscal consolidation "
        "remains a priority for many member states."
    )
    # "OECD Economic Outlook 2024" is a report title — should be italic
    # but it's in plain roman here → violates rule_140

    # ──────────────────────────────────────────────────────────────────
    # 6. BOLD unit of measure — violates rule_177
    #    Units of measure must NOT be bold or italic.
    # ──────────────────────────────────────────────────────────────────
    p6 = doc.add_paragraph()
    p6.add_run("Global temperatures have increased by approximately 1.1 ")
    run6 = p6.add_run("°C")
    run6.bold = True  # ← violation: unit of measure must not be bold
    p6.add_run(" since the pre-industrial era.")

    # ──────────────────────────────────────────────────────────────────
    # 7. ITALIC unit of measure — violates rule_177
    #    Units of measure must NOT be italicised.
    # ──────────────────────────────────────────────────────────────────
    p7 = doc.add_paragraph()
    p7.add_run("The recommended limit is 350 ")
    run7 = p7.add_run("ppm")
    run7.italic = True  # ← violation: unit of measure must not be italic
    p7.add_run(" of CO2 in the atmosphere.")

    # ──────────────────────────────────────────────────────────────────
    # 8. BOLD + ITALIC used for hierarchy — violates rule_138
    #    Styles should control hierarchy, not manual bold/italic.
    # ──────────────────────────────────────────────────────────────────
    p8 = doc.add_paragraph()
    run8 = p8.add_run("Key Findings")
    run8.bold = True   # ← violation: using bold to create sub-heading
    run8.italic = True  # ← violation: using italic to create sub-heading
    # This should be a proper heading style, not manual formatting

    # ──────────────────────────────────────────────────────────────────
    # 9. Mixed: some correct, some wrong formatting in one paragraph
    #    The foreign term "ad hoc" should be italic (correct per rule_145)
    #    but the rest of the sentence should NOT be italic.
    # ──────────────────────────────────────────────────────────────────
    p9 = doc.add_paragraph()
    run9a = p9.add_run("Several ")
    run9a.italic = True  # ← violation: plain word should not be italic
    run9b = p9.add_run("ad hoc")
    run9b.italic = True  # ← CORRECT: Latin term should be italic
    run9c = p9.add_run(" committees were established to address the crisis.")
    run9c.italic = True  # ← violation: plain text should not be italic

    # ──────────────────────────────────────────────────────────────────
    # 10. Underline + bold combo — violates rule_138, 149, 152
    # ──────────────────────────────────────────────────────────────────
    p10 = doc.add_paragraph()
    run10 = p10.add_run("Recommendations for Member Countries")
    run10.bold = True
    run10.underline = WD_UNDERLINE.SINGLE
    # Double violation: bold + underline on what should be a styled heading

    # ──────────────────────────────────────────────────────────────────
    # 11. Clean closing paragraph (no violations)
    # ──────────────────────────────────────────────────────────────────
    p11 = doc.add_paragraph()
    p11.add_run(
        "For further information, see Chapter 3, \"Fiscal Instruments for "
        "Climate Policy\", in the companion volume."
    )
    # Chapter title in roman + double quotes: CORRECT per rule_143

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Created: {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")

    # Summary of violations
    print("\n=== Embedded OpenXML Styling Violations ===")
    violations = [
        ("Para 2", "Entire body text is bold", "rule_149, rule_150"),
        ("Para 3", "Body text has underline", "rule_152"),
        ("Para 4", "'sustainability' is italicised (common word)", "rule_141"),
        ("Para 5", "'OECD Economic Outlook 2024' is NOT italic (report title)", "rule_140"),
        ("Para 6", "'°C' unit is bold", "rule_177"),
        ("Para 7", "'ppm' unit is italic", "rule_177"),
        ("Para 8", "'Key Findings' is bold+italic (hierarchy via formatting)", "rule_138, rule_150"),
        ("Para 9", "'Several' and 'committees...' runs are italic (only 'ad hoc' should be)", "rule_141"),
        ("Para 10", "'Recommendations...' is bold+underline", "rule_138, rule_149, rule_152"),
    ]
    for loc, desc, rules in violations:
        print(f"  {loc}: {desc}  [{rules}]")


if __name__ == "__main__":
    main()
